from __future__ import annotations

import argparse
import json
import os
import re
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any


CHINESE_NUMERALS = "一二三四五六七八九十百千万"
H1_PATTERN = re.compile(rf"^[{CHINESE_NUMERALS}]+、")
H2_PATTERN = re.compile(rf"^（[{CHINESE_NUMERALS}]+）")
INVALID_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]+')
ENGLISH_INLINE_SPACING = re.compile(r"(?<=[A-Za-z0-9])([,;])(?=[A-Za-z])")
CHINESE_INLINE_SPACING = re.compile(r"(?<=[，。；：？！])\s+(?=[\u3400-\u9fff])")
CONTEXT_PUNCTUATION = set(",，.。．;；:：?？！!()（）[]【】{}｛｝<>《》\"“”'‘’、")

CHINESE_PUNCTUATION = {
    ",": "，",
    "，": "，",
    ".": "。",
    "。": "。",
    "．": "。",
    ";": "；",
    "；": "；",
    ":": "：",
    "：": "：",
    "?": "？",
    "？": "？",
    "!": "！",
    "！": "！",
    "(": "（",
    "（": "（",
    ")": "）",
    "）": "）",
    "[": "【",
    "【": "【",
    "]": "】",
    "】": "】",
    "{": "｛",
    "｛": "｛",
    "}": "｝",
    "｝": "｝",
    "<": "《",
    "《": "《",
    ">": "》",
    "》": "》",
}

ENGLISH_PUNCTUATION = {
    "，": ",",
    "、": ",",
    "。": ".",
    "．": ".",
    "；": ";",
    "：": ":",
    "？": "?",
    "！": "!",
    "（": "(",
    "）": ")",
    "【": "[",
    "】": "]",
    "｛": "{",
    "｝": "}",
    "《": "<",
    "》": ">",
}

HAO_TO_PT = {
    "初号": 42.0,
    "小初": 36.0,
    "一号": 26.0,
    "小一": 24.0,
    "二号": 22.0,
    "小二": 18.0,
    "三号": 16.0,
    "小三": 15.0,
    "四号": 14.0,
    "小四": 12.0,
    "五号": 10.5,
    "小五": 9.0,
}

DEFAULT_TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "assets" / "default_template.json"


def main() -> int:
    args = parse_args()

    try:
        config = load_config(args.config)
        apply_cli_overrides(config, args)
        text = read_input_text(args)
        output_path, report = export_gongwen_docx(text, config, args.output, args.output_dir)
    except Exception as exc:
        print(f"[error] {exc}", file=sys.stderr)
        return 1

    print(f"[output] {output_path}")
    print_font_report(report)
    return 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Format Chinese official documents into .docx.")
    input_group = parser.add_mutually_exclusive_group(required=True)
    input_group.add_argument("--input", help="Input .docx or .txt path.")
    input_group.add_argument("--text-file", help="Plain text file path.")
    input_group.add_argument("--text", help="Plain text content.")

    parser.add_argument("--config", help="Optional JSON config. Missing values fall back to default_template.json.")
    parser.add_argument("--output", help="Exact output .docx path.")
    parser.add_argument("--output-dir", help="Output directory. Ignored when --output is provided.")

    parser.add_argument("--line-spacing-pt", type=float)
    parser.add_argument("--margins-mm", help="Margins as top,bottom,left,right in millimeters.")
    parser.add_argument("--no-normalize-punctuation", action="store_true")

    parser.add_argument("--title-font", help="Exact title font name. Overrides title keyword matching.")
    parser.add_argument("--title-font-keyword", help="Title font keyword to search locally.")
    parser.add_argument("--title-size")
    parser.add_argument("--h1-font")
    parser.add_argument("--h1-size")
    parser.add_argument("--h2-font")
    parser.add_argument("--h2-size")
    parser.add_argument("--body-font")
    parser.add_argument("--body-size")

    parser.add_argument("--focus-fields", help="Replace focus fields. Separate values with |, comma, or Chinese comma.")
    parser.add_argument("--append-focus-fields", help="Append focus fields. Separate values with |, comma, or Chinese comma.")
    parser.add_argument("--focus-font")
    parser.add_argument("--focus-size")
    parser.add_argument("--focus-bold", choices=["true", "false"])
    return parser.parse_args()


def load_config(path: str | None) -> dict[str, Any]:
    default_config = json.loads(DEFAULT_TEMPLATE_PATH.read_text(encoding="utf-8"))
    if not path:
        return default_config

    user_config = json.loads(Path(path).read_text(encoding="utf-8"))
    return deep_merge(default_config, user_config)


def deep_merge(base: dict[str, Any], override: dict[str, Any]) -> dict[str, Any]:
    result = deepcopy(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(result.get(key), dict):
            result[key] = deep_merge(result[key], value)
        else:
            result[key] = value
    return result


def apply_cli_overrides(config: dict[str, Any], args: argparse.Namespace) -> None:
    if args.line_spacing_pt is not None:
        config["line_spacing_pt"] = args.line_spacing_pt
    if args.margins_mm:
        values = [int(part.strip()) for part in args.margins_mm.split(",")]
        if len(values) != 4:
            raise ValueError("--margins-mm must be top,bottom,left,right")
        config["margins_mm"] = {"top": values[0], "bottom": values[1], "left": values[2], "right": values[3]}
    if args.no_normalize_punctuation:
        config["punctuation_normalization"] = False

    styles = config.setdefault("styles", {})
    set_style(styles, "title", font=args.title_font, font_keyword=args.title_font_keyword, size=args.title_size)
    set_style(styles, "h1", font=args.h1_font, size=args.h1_size)
    set_style(styles, "h2", font=args.h2_font, size=args.h2_size)
    set_style(styles, "body", font=args.body_font, size=args.body_size)

    focus = config.setdefault("focus_fields", {})
    if args.focus_fields is not None:
        focus["fields"] = parse_field_list(args.focus_fields)
    if args.append_focus_fields:
        fields = list(focus.get("fields", []))
        for value in parse_field_list(args.append_focus_fields):
            if value not in fields:
                fields.append(value)
        focus["fields"] = fields
    if args.focus_font:
        focus.setdefault("style", {})["font"] = args.focus_font
    if args.focus_size:
        focus.setdefault("style", {})["font_size"] = args.focus_size
    if args.focus_bold:
        focus["bold"] = args.focus_bold == "true"


def set_style(
    styles: dict[str, Any],
    name: str,
    *,
    font: str | None = None,
    font_keyword: str | None = None,
    size: str | None = None,
) -> None:
    style = styles.setdefault(name, {})
    if font:
        style["font"] = font
        style.pop("font_keyword", None)
    if font_keyword:
        style["font_keyword"] = font_keyword
        style.pop("font", None)
    if size:
        style["font_size"] = size


def parse_field_list(raw: str) -> list[str]:
    parts = re.split(r"[|,，、]+", raw)
    values: list[str] = []
    for part in parts:
        value = part.strip()
        if value and value not in values:
            values.append(value)
    return values


def read_input_text(args: argparse.Namespace) -> str:
    if args.text is not None:
        return args.text
    if args.text_file:
        return read_text_file(Path(args.text_file))
    if args.input:
        path = Path(args.input)
        if not path.exists():
            raise FileNotFoundError(path)
        if path.suffix.lower() == ".docx":
            return read_docx_text(path)
        return read_text_file(path)
    raise ValueError("No input provided.")


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Missing dependency python-docx. Install with: python -m pip install python-docx") from exc

    document = Document(str(path))
    return "\n".join(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())


def export_gongwen_docx(
    text: str,
    config: dict[str, Any],
    output: str | None,
    output_dir: str | None,
) -> tuple[Path, dict[str, Any]]:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError as exc:
        raise RuntimeError("Missing dependency python-docx. Install with: python -m pip install python-docx") from exc

    if config.get("punctuation_normalization", True):
        text = normalize_punctuation_for_output(text)

    paragraphs = classify_paragraphs(text)
    if not paragraphs:
        raise ValueError("No non-empty paragraphs to export.")

    installed_fonts = collect_installed_font_names()
    font_report = resolve_config_fonts(config, installed_fonts)
    target_path = build_output_path(paragraphs[0]["text"], output, output_dir)

    document = Document()
    section = document.sections[0]
    page = config.get("page", {})
    margins = config.get("margins_mm", {})
    section.page_width = Mm(page.get("width_mm", 210))
    section.page_height = Mm(page.get("height_mm", 297))
    section.top_margin = Mm(margins.get("top", 36))
    section.bottom_margin = Mm(margins.get("bottom", 36))
    section.left_margin = Mm(margins.get("left", 27))
    section.right_margin = Mm(margins.get("right", 27))

    styles = config["styles"]
    line_spacing = Pt(float(config.get("line_spacing_pt", 30)))
    alignment_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    for index, item in enumerate(paragraphs):
        paragraph_type = item["type"]
        style = styles[paragraph_type]
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = alignment_map.get(style.get("alignment"), WD_ALIGN_PARAGRAPH.LEFT)

        indent_chars = int(style.get("first_line_indent_chars", 0) or 0)
        if indent_chars:
            paragraph.paragraph_format.first_line_indent = Pt(font_size_to_pt(style.get("font_size")) * indent_chars)

        paragraph_text = item["text"].lstrip() if paragraph_type == "title" else item["text"]
        add_text_runs(paragraph, paragraph_text, style, config.get("focus_fields", {}), Pt, qn)

        if index == 0 and paragraph_type == "title" and len(paragraphs) > 1:
            body_style = styles["body"]
            blank = document.add_paragraph()
            blank.paragraph_format.line_spacing = line_spacing
            blank.paragraph_format.space_before = Pt(0)
            blank.paragraph_format.space_after = Pt(0)
            blank.paragraph_format.first_line_indent = Pt(font_size_to_pt(body_style.get("font_size")) * 2)
            blank.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    target_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(target_path))
    return target_path, font_report


def classify_paragraphs(text: str) -> list[dict[str, str]]:
    raw_paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
    classified: list[dict[str, str]] = []
    for index, paragraph in enumerate(raw_paragraphs):
        if index == 0:
            paragraph_type = "title"
        elif H1_PATTERN.match(paragraph):
            paragraph_type = "h1"
        elif H2_PATTERN.match(paragraph):
            paragraph_type = "h2"
        else:
            paragraph_type = "body"
        classified.append({"text": paragraph, "type": paragraph_type})
    return classified


def resolve_config_fonts(config: dict[str, Any], installed_fonts: list[str]) -> dict[str, Any]:
    report: dict[str, Any] = {"styles": {}, "missing_fonts": []}
    installed_set = {normalize_font_key(name): name for name in installed_fonts}

    for style_name, style in config.get("styles", {}).items():
        font = style.get("font")
        keyword = style.get("font_keyword")
        if keyword and not font:
            match = match_font_by_keyword(keyword, installed_fonts)
            selected = match["selected"] or keyword
            style["resolved_font"] = selected
            report["styles"][style_name] = {
                "mode": "keyword",
                "keyword": keyword,
                "selected": match["selected"],
                "candidates": match["candidates"],
                "used": selected,
            }
            if not match["selected"]:
                report["missing_fonts"].append({"style": style_name, "font": keyword, "reason": "keyword_not_found"})
            continue

        selected = font or ""
        style["resolved_font"] = selected
        found = normalize_font_key(selected) in installed_set if selected else False
        report["styles"][style_name] = {"mode": "exact", "requested": selected, "found": found, "used": selected}
        if selected and not found:
            report["missing_fonts"].append({"style": style_name, "font": selected, "reason": "exact_not_found"})

    focus = config.get("focus_fields", {})
    focus_style = focus.get("style", {})
    focus_font = focus_style.get("font", "")
    focus_style["resolved_font"] = focus_font
    found = normalize_font_key(focus_font) in installed_set if focus_font else False
    report["focus_fields"] = {"requested": focus_font, "found": found, "used": focus_font}
    if focus_font and not found:
        report["missing_fonts"].append({"style": "focus_fields", "font": focus_font, "reason": "exact_not_found"})

    return report


def match_font_by_keyword(keyword: str, installed_fonts: list[str]) -> dict[str, Any]:
    key = normalize_font_key(keyword)
    candidates = [name for name in installed_fonts if key in normalize_font_key(name)]
    candidates.sort(key=lambda name: font_rank(name, key))
    return {"selected": candidates[0] if candidates else None, "candidates": candidates[:10]}


def font_rank(name: str, keyword_key: str) -> tuple[int, int, int, str]:
    normalized = normalize_font_key(name)
    if normalized == keyword_key:
        primary = 0
    elif normalized.startswith(keyword_key):
        primary = 1
    elif keyword_key in normalized:
        primary = 2
    else:
        primary = 3
    return (primary, len(normalized), normalized.find(keyword_key), normalized)


def normalize_font_key(value: str) -> str:
    return re.sub(r"[\s_\-（）()]+", "", value).casefold()


def collect_installed_font_names() -> list[str]:
    names: set[str] = set()
    if os.name == "nt":
        names.update(enum_windows_gdi_font_names())
        names.update(read_windows_font_registry())
        names.update(read_windows_font_files())
        names.update(infer_windows_font_aliases(names))
    return sorted(names, key=lambda item: item.casefold())


def enum_windows_gdi_font_names() -> set[str]:
    names: set[str] = set()
    try:
        import ctypes
        from ctypes import wintypes
    except Exception:
        return names

    LF_FACESIZE = 32
    DEFAULT_CHARSET = 1

    class LOGFONTW(ctypes.Structure):
        _fields_ = [
            ("lfHeight", ctypes.c_long),
            ("lfWidth", ctypes.c_long),
            ("lfEscapement", ctypes.c_long),
            ("lfOrientation", ctypes.c_long),
            ("lfWeight", ctypes.c_long),
            ("lfItalic", ctypes.c_byte),
            ("lfUnderline", ctypes.c_byte),
            ("lfStrikeOut", ctypes.c_byte),
            ("lfCharSet", ctypes.c_byte),
            ("lfOutPrecision", ctypes.c_byte),
            ("lfClipPrecision", ctypes.c_byte),
            ("lfQuality", ctypes.c_byte),
            ("lfPitchAndFamily", ctypes.c_byte),
            ("lfFaceName", ctypes.c_wchar * LF_FACESIZE),
        ]

    FONTENUMPROCW = ctypes.WINFUNCTYPE(
        ctypes.c_int,
        ctypes.POINTER(LOGFONTW),
        wintypes.LPARAM,
        wintypes.DWORD,
        wintypes.LPARAM,
    )

    def callback(logfont: Any, _textmetric: Any, _font_type: Any, _param: Any) -> int:
        face_name = logfont.contents.lfFaceName.strip()
        if face_name:
            names.add(face_name)
        return 1

    try:
        gdi32 = ctypes.WinDLL("gdi32")
        user32 = ctypes.WinDLL("user32")
        hdc = user32.GetDC(None)
        logfont = LOGFONTW()
        logfont.lfCharSet = DEFAULT_CHARSET
        callback_func = FONTENUMPROCW(callback)
        gdi32.EnumFontFamiliesExW(hdc, ctypes.byref(logfont), callback_func, 0, 0)
        user32.ReleaseDC(None, hdc)
    except Exception:
        return names
    return names


def infer_windows_font_aliases(names: set[str]) -> set[str]:
    normalized = {normalize_font_key(name) for name in names}
    aliases: set[str] = set()
    if "simhei" in normalized:
        aliases.add("黑体")
    if "simsun" in normalized or "nsimsun" in normalized:
        aliases.add("宋体")
    if "simfang" in normalized or "fangsong" in normalized:
        aliases.add("仿宋")
    if "simkai" in normalized or "kaiti" in normalized:
        aliases.add("楷体")
    return aliases


def read_windows_font_registry() -> set[str]:
    names: set[str] = set()
    try:
        import winreg
    except ImportError:
        return names

    reg_path = r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"
    roots = [winreg.HKEY_LOCAL_MACHINE, winreg.HKEY_CURRENT_USER]
    for root in roots:
        try:
            with winreg.OpenKey(root, reg_path) as key:
                index = 0
                while True:
                    try:
                        display_name, file_value, _ = winreg.EnumValue(key, index)
                    except OSError:
                        break
                    index += 1
                    cleaned = clean_registry_font_name(str(display_name))
                    if cleaned:
                        names.add(cleaned)
                    stem = Path(str(file_value)).stem
                    if stem:
                        names.add(stem)
        except OSError:
            continue
    return names


def clean_registry_font_name(value: str) -> str:
    value = re.sub(r"\s*\([^)]*\)\s*$", "", value).strip()
    value = re.sub(r"\s+(Regular|Bold|Italic|Light|Medium|常规|粗体|斜体)$", "", value, flags=re.IGNORECASE)
    return value.strip()


def read_windows_font_files() -> set[str]:
    names: set[str] = set()
    fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    if not fonts_dir.exists():
        return names

    for path in fonts_dir.glob("*"):
        if path.suffix.lower() not in {".ttf", ".ttc", ".otf"}:
            continue
        names.add(path.stem)
        names.update(read_internal_font_names(path))
    return names


def read_internal_font_names(path: Path) -> set[str]:
    names: set[str] = set()
    try:
        from fontTools.ttLib import TTFont
    except Exception:
        return names

    try:
        font = TTFont(str(path), fontNumber=0, lazy=True)
        for record in font["name"].names:
            if record.nameID not in {1, 4, 16}:
                continue
            try:
                name = record.toUnicode().strip()
            except Exception:
                continue
            if name:
                names.add(name)
        font.close()
    except Exception:
        return names
    return names


def add_text_runs(paragraph: Any, text: str, base_style: dict[str, Any], focus_config: dict[str, Any], pt_factory: Any, qn: Any) -> None:
    fields = normalize_focus_fields(focus_config.get("fields", []))
    if not fields:
        run = paragraph.add_run(text)
        apply_run_style(run, base_style, pt_factory, qn)
        return

    for segment_text, is_focus in split_focus_segments(text, fields):
        run = paragraph.add_run(segment_text)
        if is_focus:
            focus_style = focus_config.get("style", {})
            apply_run_style(run, focus_style, pt_factory, qn, bold=bool(focus_config.get("bold", True)))
        else:
            apply_run_style(run, base_style, pt_factory, qn)


def normalize_focus_fields(values: Any) -> list[str]:
    if not isinstance(values, list):
        return []
    normalized: list[str] = []
    for value in sorted(values, key=lambda item: len(str(item)), reverse=True):
        text = str(value).strip()
        if text and text not in normalized:
            normalized.append(text)
    return normalized


def split_focus_segments(text: str, fields: list[str]) -> list[tuple[str, bool]]:
    segments: list[tuple[str, bool]] = []
    index = 0
    while index < len(text):
        matched = next((field for field in fields if text.startswith(field, index)), None)
        if matched:
            segments.append((matched, True))
            index += len(matched)
            continue

        start = index
        while index < len(text) and not any(text.startswith(field, index) for field in fields):
            index += 1
        segments.append((text[start:index], False))
    return segments


def apply_run_style(run: Any, style: dict[str, Any], pt_factory: Any, qn: Any, *, bold: bool | None = None) -> None:
    font_family = style.get("resolved_font") or style.get("font") or style.get("font_keyword") or "宋体"
    run.font.name = font_family
    run.font.size = pt_factory(font_size_to_pt(style.get("font_size")))
    if bold is not None:
        run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_family)
    r_fonts.set(qn("w:ascii"), font_family)
    r_fonts.set(qn("w:hAnsi"), font_family)


def font_size_to_pt(size_label: str | None) -> float:
    if not size_label:
        return HAO_TO_PT["小四"]
    if size_label in HAO_TO_PT:
        return HAO_TO_PT[size_label]
    try:
        return float(size_label)
    except ValueError:
        return HAO_TO_PT["小四"]


def normalize_punctuation_for_output(text: str) -> str:
    return "\n".join(normalize_line_punctuation(line) for line in text.splitlines())


def normalize_line_punctuation(text: str) -> str:
    normalized: list[str] = []
    double_quote_open = True
    single_quote_open = True

    for index, char in enumerate(text):
        if char in {'"', "“", "”"}:
            if should_use_chinese_punctuation(text, index):
                normalized.append("“" if double_quote_open else "”")
                double_quote_open = not double_quote_open
            else:
                normalized.append('"')
            continue

        if char in {"'", "‘", "’"}:
            if is_english_apostrophe(text, index):
                normalized.append("'")
            elif should_use_chinese_punctuation(text, index):
                normalized.append("‘" if single_quote_open else "’")
                single_quote_open = not single_quote_open
            else:
                normalized.append("'")
            continue

        if char in CHINESE_PUNCTUATION or char in ENGLISH_PUNCTUATION or char == "、":
            normalized.append(normalize_punctuation_char(text, index, char))
            continue

        normalized.append(char)

    normalized_text = ENGLISH_INLINE_SPACING.sub(r"\1 ", "".join(normalized))
    return CHINESE_INLINE_SPACING.sub("", normalized_text)


def normalize_punctuation_char(text: str, index: int, char: str) -> str:
    if char in {".", "。", "．"} and is_numeric_or_ordered_list_period(text, index):
        return "."
    if should_use_chinese_punctuation(text, index):
        return CHINESE_PUNCTUATION.get(char, char)
    return ENGLISH_PUNCTUATION.get(char, char)


def should_use_chinese_punctuation(text: str, index: int) -> bool:
    left = nearest_context_char(text, index, -1)
    right = nearest_context_char(text, index, 1)
    left_kind = context_kind(left)
    right_kind = context_kind(right)

    if left_kind == "english" and right_kind == "english":
        return False
    if right_kind == "english" and left_kind is None:
        return False
    if left_kind == "english" and right_kind is None:
        return has_chinese_sentence_context(text, index)
    if left_kind == "chinese" or right_kind == "chinese":
        return True
    return False


def nearest_context_char(text: str, index: int, step: int) -> str | None:
    cursor = index + step
    while 0 <= cursor < len(text):
        char = text[cursor]
        if char.isspace() or char in CONTEXT_PUNCTUATION:
            cursor += step
            continue
        return char
    return None


def context_kind(char: str | None) -> str | None:
    if char is None:
        return None
    if is_chinese_char(char):
        return "chinese"
    if char.isascii() and (char.isalpha() or char.isdigit()):
        return "english"
    return None


def has_chinese_sentence_context(text: str, index: int) -> bool:
    cursor = index - 1
    while cursor >= 0:
        char = text[cursor]
        if char in "。.!?！？;；":
            return False
        if is_chinese_char(char):
            return True
        cursor -= 1
    return False


def is_chinese_char(char: str) -> bool:
    return "\u3400" <= char <= "\u4dbf" or "\u4e00" <= char <= "\u9fff" or "\uf900" <= char <= "\ufaff"


def is_english_apostrophe(text: str, index: int) -> bool:
    left = text[index - 1] if index > 0 else ""
    right = text[index + 1] if index + 1 < len(text) else ""
    return left.isascii() and right.isascii() and left.isalpha() and right.isalpha()


def is_numeric_or_ordered_list_period(text: str, index: int) -> bool:
    left = nearest_context_char(text, index, -1)
    right = nearest_context_char(text, index, 1)
    if left and right and left.isdigit() and right.isdigit():
        return True
    prefix = text[:index].strip()
    suffix = text[index + 1 :]
    return bool(prefix and len(prefix) <= 3 and prefix.isdigit() and suffix[:1].isspace())


def build_output_path(title: str, output: str | None, output_dir: str | None) -> Path:
    if output:
        return Path(output).expanduser()

    directory = Path(output_dir).expanduser() if output_dir else Path.cwd()
    safe_title = sanitize_filename(title) or "排版结果"
    base_name = f"{safe_title}_排版结果"
    target = directory / f"{base_name}.docx"
    suffix_index = 1
    while target.exists():
        target = directory / f"{base_name}{suffix_index}.docx"
        suffix_index += 1
    return target


def sanitize_filename(value: str) -> str:
    return INVALID_FILENAME_CHARS.sub("_", value).strip().strip(".")


def print_font_report(report: dict[str, Any]) -> None:
    for style_name, info in report.get("styles", {}).items():
        if info.get("mode") == "keyword":
            selected = info.get("selected")
            if selected:
                print(f"[font] {style_name}: keyword \"{info['keyword']}\" matched \"{selected}\".")
            else:
                print(f"[font-warning] {style_name}: keyword \"{info['keyword']}\" was not found locally; using requested name.")
            candidates = info.get("candidates") or []
            if candidates:
                print(f"[font-candidates] {style_name}: {', '.join(candidates[:5])}")
        else:
            status = "found" if info.get("found") else "not found"
            print(f"[font] {style_name}: \"{info.get('requested')}\" {status}.")

    focus = report.get("focus_fields", {})
    if focus:
        status = "found" if focus.get("found") else "not found"
        print(f"[font] focus_fields: \"{focus.get('requested')}\" {status}.")

    missing = report.get("missing_fonts") or []
    if missing:
        names = ", ".join(f"{item['style']}={item['font']}" for item in missing)
        print(f"[font-warning] Missing or unmatched fonts: {names}. Install them or specify replacements if exact typography is required.")


if __name__ == "__main__":
    raise SystemExit(main())
